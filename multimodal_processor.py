#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Multimodal Input Processor
===========================

Procesa entradas multimodales: audio, imágenes y documentos.
Integración con WhatsApp Business para agentes de ventas.
"""

import os
import base64
from pathlib import Path
from typing import Dict, Any, Optional, List, Union
from dataclasses import dataclass
import json


@dataclass
class MultimodalInput:
    """Contexto unificado para entrada multimodal"""
    input_type: str  # 'text', 'audio', 'image', 'document'
    content: str  # Texto extraído o descripción
    metadata: Dict[str, Any]
    original_file: Optional[str] = None
    confidence: float = 1.0
    processing_notes: List[str] = None


class MultimodalProcessor:
    """Procesador de entradas multimodales"""
    
    def __init__(self, openai_client=None):
        """
        Inicializa el procesador multimodal
        
        Args:
            openai_client: Cliente de OpenAI para Whisper y Vision
        """
        self.openai_client = openai_client
        self.supported_audio_formats = ['.mp3', '.wav', '.m4a', '.ogg', '.webm']
        self.supported_image_formats = ['.jpg', '.jpeg', '.png', '.gif', '.webp']
        self.supported_doc_formats = ['.pdf', '.docx', '.txt']
        
    def process_input(self, input_data: Union[str, bytes, Path], 
                     input_type: Optional[str] = None) -> MultimodalInput:
        """
        Procesa entrada multimodal y la convierte en contexto unificado
        
        Args:
            input_data: Datos de entrada (ruta, bytes, o texto)
            input_type: Tipo de entrada ('auto' para detección automática)
            
        Returns:
            MultimodalInput con el contexto procesado
        """
        # Detectar tipo automáticamente si no se especifica
        if input_type is None or input_type == 'auto':
            input_type = self._detect_input_type(input_data)
        
        processing_notes = []
        
        try:
            if input_type == 'text':
                return self._process_text(input_data)
            elif input_type == 'audio':
                return self._process_audio(input_data, processing_notes)
            elif input_type == 'image':
                return self._process_image(input_data, processing_notes)
            elif input_type == 'document':
                return self._process_document(input_data, processing_notes)
            else:
                raise ValueError(f"Tipo de entrada no soportado: {input_type}")
        except Exception as e:
            processing_notes.append(f"Error procesando: {str(e)}")
            return MultimodalInput(
                input_type='error',
                content=f"Error al procesar entrada: {str(e)}",
                metadata={'error': str(e)},
                confidence=0.0,
                processing_notes=processing_notes
            )
    
    def _detect_input_type(self, input_data: Union[str, bytes, Path]) -> str:
        """Detecta automáticamente el tipo de entrada"""
        if isinstance(input_data, str) and not os.path.exists(input_data):
            return 'text'
        
        # Si es una ruta de archivo
        if isinstance(input_data, (str, Path)):
            file_path = Path(input_data)
            if file_path.exists():
                ext = file_path.suffix.lower()
                if ext in self.supported_audio_formats:
                    return 'audio'
                elif ext in self.supported_image_formats:
                    return 'image'
                elif ext in self.supported_doc_formats:
                    return 'document'
        
        # Si son bytes, intentar detectar por contenido
        if isinstance(input_data, bytes):
            # Simple heuristic: check magic numbers
            if input_data.startswith(b'\xff\xd8\xff'):  # JPEG
                return 'image'
            elif input_data.startswith(b'\x89PNG'):  # PNG
                return 'image'
            elif input_data.startswith(b'%PDF'):  # PDF
                return 'document'
        
        return 'text'
    
    def _process_text(self, text: str) -> MultimodalInput:
        """Procesa entrada de texto simple"""
        return MultimodalInput(
            input_type='text',
            content=str(text),
            metadata={'length': len(str(text))},
            confidence=1.0,
            processing_notes=[]
        )
    
    def _process_audio(self, audio_data: Union[str, bytes, Path], 
                      notes: List[str]) -> MultimodalInput:
        """
        Procesa audio usando Whisper API
        
        Args:
            audio_data: Datos de audio (ruta o bytes)
            notes: Lista para agregar notas de procesamiento
            
        Returns:
            MultimodalInput con transcripción
        """
        notes.append("Procesando audio con Whisper API")
        
        if not self.openai_client:
            notes.append("Cliente OpenAI no disponible, retornando placeholder")
            return MultimodalInput(
                input_type='audio',
                content="[Audio no transcrito: cliente OpenAI no configurado]",
                metadata={'processed': False},
                confidence=0.0,
                processing_notes=notes
            )
        
        try:
            # Obtener ruta del archivo
            audio_path = None
            if isinstance(audio_data, bytes):
                # Guardar bytes temporalmente
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.webm', delete=False) as f:
                    f.write(audio_data)
                    audio_path = f.name
            else:
                audio_path = str(audio_data)
            
            try:
                # Transcribir con Whisper
                with open(audio_path, 'rb') as audio_file:
                    transcription = self.openai_client.audio.transcriptions.create(
                        model="whisper-1",
                        file=audio_file,
                        language="es"
                    )
                
                notes.append("Transcripción completada exitosamente")
                
                return MultimodalInput(
                    input_type='audio',
                    content=transcription.text,
                    metadata={
                        'transcription_model': 'whisper-1',
                        'language': 'es'
                    },
                    original_file=str(audio_data) if not isinstance(audio_data, bytes) else None,
                    confidence=0.9,
                    processing_notes=notes
                )
            finally:
                # Limpiar archivo temporal si se creó
                if isinstance(audio_data, bytes) and audio_path and os.path.exists(audio_path):
                    try:
                        os.unlink(audio_path)
                    except Exception as e:
                        notes.append(f"No se pudo eliminar archivo temporal: {str(e)}")
        
        except Exception as e:
            notes.append(f"Error en transcripción: {str(e)}")
            return MultimodalInput(
                input_type='audio',
                content=f"[Error transcribiendo audio: {str(e)}]",
                metadata={'error': str(e)},
                confidence=0.0,
                processing_notes=notes
            )
    
    def _process_image(self, image_data: Union[str, bytes, Path], 
                      notes: List[str]) -> MultimodalInput:
        """
        Procesa imagen usando Vision API
        
        Args:
            image_data: Datos de imagen (ruta o bytes)
            notes: Lista para agregar notas de procesamiento
            
        Returns:
            MultimodalInput con descripción de la imagen
        """
        notes.append("Procesando imagen con Vision API")
        
        if not self.openai_client:
            notes.append("Cliente OpenAI no disponible, procesando con PIL básico")
            return self._process_image_basic(image_data, notes)
        
        try:
            # Convertir imagen a base64 si es necesario
            if isinstance(image_data, bytes):
                image_base64 = base64.b64encode(image_data).decode('utf-8')
            else:
                with open(str(image_data), 'rb') as f:
                    image_base64 = base64.b64encode(f.read()).decode('utf-8')
            
            # Usar GPT-4 Vision para analizar
            response = self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Describe esta imagen en español. Si es una foto de un producto de construcción, describe sus características técnicas visibles."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{image_base64}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=300
            )
            
            description = response.choices[0].message.content
            notes.append("Análisis de imagen completado")
            
            return MultimodalInput(
                input_type='image',
                content=description,
                metadata={
                    'vision_model': 'gpt-4o',
                    'analysis_type': 'product_description'
                },
                original_file=str(image_data) if not isinstance(image_data, bytes) else None,
                confidence=0.85,
                processing_notes=notes
            )
        
        except Exception as e:
            notes.append(f"Error en análisis de imagen: {str(e)}")
            return self._process_image_basic(image_data, notes)
    
    def _process_image_basic(self, image_data: Union[str, bytes, Path], 
                            notes: List[str]) -> MultimodalInput:
        """Procesamiento básico de imagen sin Vision API"""
        try:
            from PIL import Image
            import io
            
            if isinstance(image_data, bytes):
                img = Image.open(io.BytesIO(image_data))
            else:
                img = Image.open(str(image_data))
            
            width, height = img.size
            format_name = img.format or 'unknown'
            
            notes.append(f"Imagen procesada: {width}x{height} {format_name}")
            
            return MultimodalInput(
                input_type='image',
                content=f"[Imagen {width}x{height} formato {format_name}. Análisis visual no disponible]",
                metadata={
                    'width': width,
                    'height': height,
                    'format': format_name
                },
                original_file=str(image_data) if not isinstance(image_data, bytes) else None,
                confidence=0.3,
                processing_notes=notes
            )
        
        except Exception as e:
            notes.append(f"Error en procesamiento básico: {str(e)}")
            return MultimodalInput(
                input_type='image',
                content="[Imagen no procesable]",
                metadata={'error': str(e)},
                confidence=0.0,
                processing_notes=notes
            )
    
    def _process_document(self, doc_data: Union[str, bytes, Path], 
                         notes: List[str]) -> MultimodalInput:
        """
        Extrae texto de documentos (PDF, DOCX, TXT)
        
        Args:
            doc_data: Datos del documento
            notes: Lista para agregar notas de procesamiento
            
        Returns:
            MultimodalInput con texto extraído
        """
        notes.append("Procesando documento")
        
        try:
            file_path = Path(doc_data) if not isinstance(doc_data, bytes) else None
            ext = file_path.suffix.lower() if file_path else '.pdf'
            
            if ext == '.txt':
                if isinstance(doc_data, bytes):
                    text = doc_data.decode('utf-8')
                else:
                    with open(str(doc_data), 'r', encoding='utf-8') as f:
                        text = f.read()
            
            elif ext == '.pdf':
                text = self._extract_pdf_text(doc_data, notes)
            
            elif ext == '.docx':
                text = self._extract_docx_text(doc_data, notes)
            
            else:
                raise ValueError(f"Formato de documento no soportado: {ext}")
            
            notes.append(f"Texto extraído: {len(text)} caracteres")
            
            return MultimodalInput(
                input_type='document',
                content=text,
                metadata={
                    'format': ext,
                    'length': len(text)
                },
                original_file=str(doc_data) if file_path else None,
                confidence=0.95,
                processing_notes=notes
            )
        
        except Exception as e:
            notes.append(f"Error extrayendo texto: {str(e)}")
            return MultimodalInput(
                input_type='document',
                content=f"[Error procesando documento: {str(e)}]",
                metadata={'error': str(e)},
                confidence=0.0,
                processing_notes=notes
            )
    
    def _extract_pdf_text(self, pdf_data: Union[str, bytes, Path], 
                         notes: List[str]) -> str:
        """Extrae texto de PDF"""
        try:
            from PyPDF2 import PdfReader
            import io
            
            if isinstance(pdf_data, bytes):
                pdf_file = io.BytesIO(pdf_data)
            else:
                pdf_file = open(str(pdf_data), 'rb')
            
            reader = PdfReader(pdf_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            if not isinstance(pdf_data, bytes):
                pdf_file.close()
            
            notes.append(f"PDF: {len(reader.pages)} páginas")
            return text.strip()
        
        except Exception as e:
            notes.append(f"Error PyPDF2: {str(e)}")
            return f"[No se pudo extraer texto del PDF: {str(e)}]"
    
    def _extract_docx_text(self, docx_data: Union[str, bytes, Path], 
                          notes: List[str]) -> str:
        """Extrae texto de DOCX"""
        try:
            from docx import Document
            import io
            
            if isinstance(docx_data, bytes):
                doc = Document(io.BytesIO(docx_data))
            else:
                doc = Document(str(docx_data))
            
            text = "\n".join([para.text for para in doc.paragraphs])
            
            notes.append(f"DOCX: {len(doc.paragraphs)} párrafos")
            return text.strip()
        
        except Exception as e:
            notes.append(f"Error python-docx: {str(e)}")
            return f"[No se pudo extraer texto del DOCX: {str(e)}]"


# Función auxiliar para crear procesador con cliente OpenAI
def create_multimodal_processor(openai_api_key: Optional[str] = None) -> MultimodalProcessor:
    """
    Crea un procesador multimodal configurado
    
    Args:
        openai_api_key: API key de OpenAI (opcional, se lee de env si no se provee)
        
    Returns:
        MultimodalProcessor configurado
    """
    openai_client = None
    
    if openai_api_key or os.getenv('OPENAI_API_KEY'):
        try:
            from openai import OpenAI
            openai_client = OpenAI(api_key=openai_api_key or os.getenv('OPENAI_API_KEY'))
        except Exception as e:
            print(f"⚠️  No se pudo inicializar cliente OpenAI: {e}")
    
    return MultimodalProcessor(openai_client=openai_client)
